from toolbox import update_ui
from toolbox import CatchException, report_exception
from toolbox import write_history_to_file, promote_file_to_downloadzone
from crazy_functions.crazy_utils import request_gpt_model_in_new_thread_with_ui_alive

import os
import sys
import re
import copy
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

_CURRENT_DIR = os.path.dirname(os.path.abspath(__file__)) if "__file__" in globals() else os.getcwd()
if _CURRENT_DIR not in sys.path:
    sys.path.insert(0, _CURRENT_DIR)

try:
    from crazy_functions.word_plugin_utils import (
        strip_model_noise,
        safe_load_json,
        now_str,
        sanitize_filename,
        read_docx_content,
        paragraph_issue_heuristics,
        chunk_text,
        estimate_char_budget,
        mark_docx_paragraphs,
    )
except ImportError:
    from word_plugin_utils import (
        strip_model_noise,
        safe_load_json,
        now_str,
        sanitize_filename,
        read_docx_content,
        paragraph_issue_heuristics,
        chunk_text,
        estimate_char_budget,
        mark_docx_paragraphs,
    )

fast_debug = False


def _strip_protocol_noise(text: str) -> str:
    text = text or ""
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.S | re.I)
    text = re.sub(r"</?json>", "", text, flags=re.I)
    text = re.sub(r"```(?:json|text|markdown)?\s*", "", text, flags=re.I)
    text = text.replace("```", "")
    return text.strip()


def _clean_obj(obj):
    if isinstance(obj, dict):
        return {k: _clean_obj(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_clean_obj(v) for v in obj]
    if isinstance(obj, str):
        return _strip_protocol_noise(strip_model_noise(obj)).strip(" \n\r\t,")
    return obj


def _extract_json_field_text(text: str, keys):
    raw = text or ""
    for key in keys:
        m = re.search(rf'"{re.escape(key)}"\s*:\s*"((?:\\.|[^"\\])*)"', raw, flags=re.S)
        if m:
            s = m.group(1)
            s = s.replace('\\n', '\n').replace('\\t', '\t').replace('\\"', '"')
            return _strip_protocol_noise(s).strip()
    return ""


def _best_effort_text(text: str, preferred_keys, max_len: int) -> str:
    obj = safe_load_json(text, None)
    if isinstance(obj, dict):
        obj = _clean_obj(obj)
        for key in preferred_keys:
            value = obj.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()[:max_len]
    field_text = _extract_json_field_text(text, preferred_keys)
    if field_text:
        return field_text[:max_len]
    cleaned = _strip_protocol_noise(strip_model_noise(text))
    lines = []
    for line in cleaned.splitlines():
        s = line.strip().strip(",")
        if not s or s in ["{", "}", "[", "]"]:
            continue
        if re.match(r'^"[^\"]+"\s*:', s):
            continue
        lines.append(s)
    return " ".join(lines).strip()[:max_len]


def _safe_json_data(text: str, default: dict, preferred_keys, max_len: int) -> dict:
    data = safe_load_json(text, None)
    if not isinstance(data, dict):
        data = copy.deepcopy(default)
    else:
        data = _clean_obj(data)
        merged = copy.deepcopy(default)
        merged.update(data)
        data = merged
    fallback = _best_effort_text(text, preferred_keys=preferred_keys, max_len=max_len)
    for key in preferred_keys:
        if key in data and not str(data.get(key, "")).strip() and fallback:
            data[key] = fallback
    return data


def _ensure_list(value):
    if value is None:
        return []
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    if isinstance(value, str):
        txt = _strip_protocol_noise(strip_model_noise(value)).strip()
        if not txt:
            return []
        parts = re.split(r'[\n；;]+', txt)
        return [p.strip(' -•') for p in parts if p.strip(' -•')]
    return [str(value).strip()]


def _clip_excerpt(text: str, limit: int = 50) -> str:
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    if len(text) <= limit:
        return text
    return text[:limit] + "…"


def _set_table_widths(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for i, width in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(width)


def _set_cell_text(cell, text):
    cell.text = str(text or "")
    for p in cell.paragraphs:
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0


def export_revision_report(doc_title, review_mode, candidate_count, issue_count, final_data, save_dir="./tmp"):
    import os
    import time
    from pathlib import Path
    from docx import Document

    Path(save_dir).mkdir(parents=True, exist_ok=True)
    out_path = os.path.join(save_dir, f"文本修改建议报告_{sanitize_filename(doc_title)}_{int(time.time())}.docx")

    doc = Document()
    doc.add_heading(f"文本修改建议报告 - {doc_title}", level=0)
    doc.add_paragraph(f"生成时间：{now_str()}")
    doc.add_paragraph(f"审阅模式：{review_mode}")
    doc.add_paragraph(f"规则初筛候选段落数：{candidate_count}")
    doc.add_paragraph(f"最终识别问题数：{issue_count}")

    overall_findings = _ensure_list(final_data.get("overall_findings", []))
    if overall_findings:
        doc.add_heading("全文关键问题总览", level=1)
        for item in overall_findings:
            doc.add_paragraph(str(item), style="List Bullet")

    issues = final_data.get("top_issues", [])
    if issues:
        doc.add_heading("关键问题定位表", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        _set_table_widths(table, [2.0, 1.6, 2.2, 8.8, 5.0])

        hdr = table.rows[0].cells
        _set_cell_text(hdr[0], "段落ID")
        _set_cell_text(hdr[1], "严重度")
        _set_cell_text(hdr[2], "问题类型")
        _set_cell_text(hdr[3], "原文摘录（前50字）")
        _set_cell_text(hdr[4], "修改建议")

        for item in issues:
            row = table.add_row().cells
            _set_cell_text(row[0], item.get("paragraph_id", ""))
            _set_cell_text(row[1], item.get("severity", ""))
            _set_cell_text(row[2], item.get("problem_type", ""))
            _set_cell_text(row[3], _clip_excerpt(item.get("original_excerpt", ""), 50))
            _set_cell_text(row[4], item.get("suggestion", ""))

        doc.add_heading("重点段落修订稿", level=1)
        for idx, item in enumerate(issues, 1):
            doc.add_heading(f"问题 {idx} - {item.get('paragraph_id', '')}", level=2)
            doc.add_paragraph(f"严重度：{item.get('severity', '')}")
            doc.add_paragraph(f"问题类型：{item.get('problem_type', '')}")
            doc.add_paragraph(f"原文：{item.get('original_excerpt', '')}")
            doc.add_paragraph(f"建议：{item.get('suggestion', '')}")
            doc.add_paragraph(f"修订稿：{item.get('rewrite', '')}")
    else:
        doc.add_heading("审阅结果", level=1)
        doc.add_paragraph("未识别出需要优先修改的关键问题。")

    doc.save(out_path)
    return out_path


def _build_review_prompt(rel_path, chunk_idx, chunk_total, chunk_text, style_hint, domain_hint, audience_hint):
    prefs = []
    if style_hint:
        prefs.append(f"目标风格：{style_hint}")
    if domain_hint:
        prefs.append(f"领域背景：{domain_hint}")
    if audience_hint:
        prefs.append(f"预期读者：{audience_hint}")
    pref_text = "；".join(prefs) if prefs else "无特别偏好"

    return (
        "你是一名严谨的学术与技术写作编辑。请审阅以下候选段落，只保留真正值得修改的关键问题。"
        "不要输出思考过程，不要写铺垫。\n"
        f"编辑偏好：{pref_text}\n"
        "请严格仅输出 <json>...</json>。\n"
        "输出格式：\n"
        "<json>{\n"
        '  "chunk_summary": "120到180字的片段审阅结论，需要说明该批段落最突出的共性问题",\n'
        '  "issues": [\n'
        '    {"paragraph_id": "P0001", "severity": "高/中/低", "problem_type": "逻辑/结构/冗长/术语/语法/格式", "original_excerpt": "原文片段", "suggestion": "一句话建议", "rewrite": "优化后的写法"}\n'
        '  ]\n'
        "}</json>\n\n"
        f"文件：{rel_path}；候选片段批次：{chunk_idx}/{chunk_total}\n"
        "候选内容如下：\n"
        f"```text\n{chunk_text}\n```"
    )


def _build_light_merge_prompt(doc_name, all_issues, max_keep):
    import json
    return f"""你是一名中文写作总编。请基于下面的问题清单做轻量汇总。不要输出思考过程。
请严格仅输出 <json>...</json>。
目标：
1. 对跨片段重复问题去重；
2. 按修改优先级筛出最值得优先处理的问题；
3. 给出 3 到 6 条全文层面的关键问题总览。
要求：
1. 不要生成执行摘要；
2. 不要生成全局统一建议；
3. top_issues 最多保留 {max_keep} 条；
4. 内容必须贴合给定问题清单，不能编造；
5. 输出语言简洁、正式，便于直接写入报告。
输出格式：
<json>{{
  "overall_findings": ["总览1", "总览2", "总览3"],
  "top_issues": [
    {{"paragraph_id": "P0001", "severity": "高/中/低", "problem_type": "逻辑/结构/冗长/术语/语法/格式", "original_excerpt": "原文片段", "suggestion": "一句话建议", "rewrite": "优化后的写法"}}
  ]
}}</json>

文档名称：{doc_name}
问题清单如下：
```json
{json.dumps(all_issues, ensure_ascii=False, indent=2)}
```"""


def _pack_candidate_text(candidates):
    lines = []
    for item in candidates:
        lines.append(f"[{item['paragraph_id']}] 规则初筛：{'；'.join(item['reasons'])}")
        lines.append(item["text"])
        lines.append("")
    return "\n".join(lines)


def _dedup_and_rank_issues(all_issues):
    severity_rank = {"高": 3, "中": 2, "低": 1}
    dedup = {}
    for item in all_issues:
        pid = item.get("paragraph_id", "")
        if not pid:
            continue
        old = dedup.get(pid)
        new_rank = severity_rank.get(item.get("severity", "中"), 2)
        old_rank = severity_rank.get(old.get("severity", "中"), 2) if old else -1
        if old is None or new_rank > old_rank:
            dedup[pid] = item
    items = list(dedup.values())
    items.sort(key=lambda x: (-severity_rank.get(x.get("severity", "中"), 2), x.get("paragraph_id", "")))
    return items


def 解析word文本建议(file_manifest, project_folder, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt):
    import os
    import json

    style_hint = plugin_kwargs.get("style", "").strip() if isinstance(plugin_kwargs, dict) else ""
    domain_hint = plugin_kwargs.get("domain", "").strip() if isinstance(plugin_kwargs, dict) else ""
    audience_hint = plugin_kwargs.get("audience", "").strip() if isinstance(plugin_kwargs, dict) else ""
    max_review_paragraphs = int(plugin_kwargs.get("max_review_paragraphs", 30)) if isinstance(plugin_kwargs, dict) else 30
    max_mark_issues = int(plugin_kwargs.get("max_mark_issues", 20)) if isinstance(plugin_kwargs, dict) else 20
    merge_issue_threshold = int(plugin_kwargs.get("merge_issue_threshold", 12)) if isinstance(plugin_kwargs, dict) else 12

    for fp in file_manifest:
        if not fp.lower().endswith(".docx"):
            raise RuntimeError("优化版当前建议仅处理 .docx 文件；.doc 请先转换为 .docx。")

        doc, paragraphs, tables = read_docx_content(fp, include_tables=False)
        char_budget = estimate_char_budget(llm_kwargs, default_budget=4500)
        candidates = paragraph_issue_heuristics(paragraphs, max_candidates=max_review_paragraphs)

        if not candidates:
            meaningful = [p for p in paragraphs if p.get("text")]
            candidates = [{
                "paragraph_id": p["id"],
                "score": 1,
                "reasons": ["常规抽样检查"],
                "excerpt": p["text"][:180],
                "text": p["text"],
            } for p in meaningful[:min(10, len(meaningful))]]
            review_mode = "抽样审阅模式"
        else:
            review_mode = "关键段落精审模式"

        candidate_text = _pack_candidate_text(candidates)
        chunks = chunk_text(candidate_text, max_chars=char_budget)
        if not chunks:
            chunks = [candidate_text[:char_budget]]

        rel_path = os.path.relpath(fp, project_folder)
        all_issues = []

        for i, chunk in enumerate(chunks, 1):
            i_say = _build_review_prompt(rel_path, i, len(chunks), chunk, style_hint, domain_hint, audience_hint)
            i_say_show_user = f"正在生成文本修改建议：{os.path.abspath(fp)} 第 {i}/{len(chunks)} 批候选段落"
            gpt_say = yield from request_gpt_model_in_new_thread_with_ui_alive(
                inputs=i_say,
                inputs_show_user=i_say_show_user,
                llm_kwargs=llm_kwargs,
                chatbot=chatbot,
                history=[],
                sys_prompt="文本修改建议，只输出结论。"
            )
            data = _safe_json_data(
                gpt_say,
                default={"chunk_summary": "", "issues": []},
                preferred_keys=["chunk_summary", "summary"],
                max_len=260,
            )
            clean_view = (data.get("chunk_summary") or _best_effort_text(gpt_say, ["chunk_summary", "summary"], 180) or "（片段审阅结论解析失败）")
            chatbot[-1] = (i_say_show_user, clean_view)
            history.extend([i_say_show_user, clean_view])
            all_issues.extend(data.get("issues", []))

        all_issues = _dedup_and_rank_issues(all_issues)

        need_merge = len(chunks) > 1 or len(all_issues) > merge_issue_threshold
        if need_merge:
            merge_prompt = _build_light_merge_prompt(os.path.basename(fp), all_issues, max_mark_issues)
            merge_show = f"正在汇总关键问题：{os.path.abspath(fp)}"
            gpt_final = yield from request_gpt_model_in_new_thread_with_ui_alive(
                inputs=merge_prompt,
                inputs_show_user=merge_show,
                llm_kwargs=llm_kwargs,
                chatbot=chatbot,
                history=[],
                sys_prompt="文本修改建议，只输出结论。"
            )
            final_data = _safe_json_data(
                gpt_final,
                default={"overall_findings": [], "top_issues": all_issues[:max_mark_issues]},
                preferred_keys=[],
                max_len=800,
            )
            if not final_data.get("top_issues"):
                final_data["top_issues"] = all_issues[:max_mark_issues]
            overview = _ensure_list(final_data.get("overall_findings", []))
            clean_final = "；".join(overview[:3]) if overview else f"已汇总 {len(final_data.get('top_issues', []))} 条关键问题"
            chatbot[-1] = (merge_show, clean_final)
            history.extend([merge_show, clean_final])
        else:
            final_data = {
                "overall_findings": [],
                "top_issues": all_issues[:max_mark_issues],
            }

        final_issues = _dedup_and_rank_issues(final_data.get("top_issues", []))[:max_mark_issues]
        final_data["top_issues"] = final_issues

        report_path = export_revision_report(
            doc_title=os.path.basename(fp),
            review_mode=review_mode,
            candidate_count=len(candidates),
            issue_count=len(final_issues),
            final_data=final_data,
            save_dir="./tmp"
        )
        json_path = report_path.replace(".docx", ".json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump({
                "file": fp,
                "review_mode": review_mode,
                "candidate_count": len(candidates),
                "all_issues": all_issues,
                "final_data": final_data,
            }, f, ensure_ascii=False, indent=2)

        marked_path = report_path.replace(".docx", "_原文标注版.docx")
        try:
            mark_docx_paragraphs(fp, final_issues[:max_mark_issues], marked_path)
        except Exception as e:
            marked_path = ""
            chatbot.append(("⚠️ 原文标注失败", str(e)))

        promote_file_to_downloadzone(report_path, chatbot=chatbot)
        promote_file_to_downloadzone(json_path, chatbot=chatbot)
        if marked_path:
            promote_file_to_downloadzone(marked_path, chatbot=chatbot)
        chatbot.append(("📄 已生成文本修改建议报告", report_path))
        if marked_path:
            chatbot.append(("🖍️ 已生成原文标注版", marked_path))
        yield from update_ui(chatbot=chatbot, history=history)

    res = write_history_to_file(history)
    promote_file_to_downloadzone(res, chatbot=chatbot)
    chatbot.append(("全部文本修改建议任务完成", res))
    yield from update_ui(chatbot=chatbot, history=history)


@CatchException
def 文本修改建议(txt, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt, user_request):
    import glob
    import os

    chatbot.append([
        "函数插件功能？",
        "批量读取 Word 文档并生成文本修改建议（轻量汇总版）：不生成执行摘要和全局统一建议，仅保留关键问题定位、修订建议与原文标注。"
    ])
    yield from update_ui(chatbot=chatbot, history=history)

    try:
        from docx import Document  # noqa: F401
    except Exception:
        report_exception(
            chatbot, history,
            a=f"解析项目: {txt}",
            b="缺少依赖，请安装：```pip install --upgrade python-docx pywin32```。"
        )
        yield from update_ui(chatbot=chatbot, history=history)
        return

    history = []

    if os.path.exists(txt):
        project_folder = txt
    else:
        if txt == "":
            txt = '空输入'
        report_exception(chatbot, history, a=f"解析项目: {txt}", b=f"找不到路径或无权限访问: {txt}")
        yield from update_ui(chatbot=chatbot, history=history)
        return

    if txt.lower().endswith('.docx'):
        file_manifest = [txt]
        project_folder = os.path.dirname(os.path.abspath(txt)) or os.getcwd()
    else:
        project_folder = os.path.abspath(project_folder)
        file_manifest = [f for f in glob.glob(f'{project_folder}/**/*.docx', recursive=True)]

    if len(file_manifest) == 0:
        report_exception(chatbot, history, a=f"解析项目: {txt}", b="未找到任何 Word 文件。")
        yield from update_ui(chatbot=chatbot, history=history)
        return

    yield from 解析word文本建议(file_manifest, project_folder, llm_kwargs, plugin_kwargs or {}, chatbot, history, system_prompt)
