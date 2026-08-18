import os
import re
import json
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Tuple, Any


def strip_model_noise(text: str) -> str:
    """清理 deepseek 等模型可能输出的思考过程与代码围栏。"""
    text = text or ""
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.S | re.I)
    text = re.sub(r"```json\s*", "", text, flags=re.I)
    text = re.sub(r"```text\s*", "", text, flags=re.I)
    text = re.sub(r"```markdown\s*", "", text, flags=re.I)
    text = text.replace("```", "")
    text = text.strip()
    return text


def extract_tagged_json(text: str) -> str:
    """优先抽取 <json>...</json> 中的内容；否则尝试寻找最外层 JSON。"""
    cleaned = strip_model_noise(text)
    m = re.search(r"<json>\s*(\{.*\}|\[.*\])\s*</json>", cleaned, flags=re.S | re.I)
    if m:
        return m.group(1).strip()

    # 兜底：提取第一个 { 或 [ 开始的 JSON 片段
    start_obj = cleaned.find("{")
    start_arr = cleaned.find("[")
    starts = [x for x in [start_obj, start_arr] if x != -1]
    if not starts:
        return ""
    start = min(starts)
    candidate = cleaned[start:]

    # 简单括号平衡截断
    stack = []
    in_string = False
    escape = False
    for i, ch in enumerate(candidate):
        if escape:
            escape = False
            continue
        if ch == "\\":
            escape = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch in "[{":
            stack.append(ch)
        elif ch in "]}":
            if stack:
                stack.pop()
            if not stack:
                return candidate[: i + 1]
    return ""



def safe_load_json(text: str, default: Any):
    payload = extract_tagged_json(text)
    if not payload:
        return default
    try:
        return json.loads(payload)
    except Exception:
        # 尝试去除尾随逗号
        payload = re.sub(r",\s*([}\]])", r"\1", payload)
        try:
            return json.loads(payload)
        except Exception:
            return default



def now_str() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")



def sanitize_filename(name: str) -> str:
    name = re.sub(r"[\\/:*?\"<>|]", "_", name)
    return name[:120] if name else "output"



def read_docx_content(fp: str, include_tables: bool = True):
    from docx import Document

    doc = Document(fp)
    paragraphs = []
    for idx, para in enumerate(doc.paragraphs, 1):
        text = (para.text or "").strip()
        style = getattr(para.style, "name", "") if para.style else ""
        paragraphs.append({
            "id": f"P{idx:04d}",
            "index": idx - 1,
            "style": style,
            "text": text,
        })

    tables = []
    if include_tables:
        for t_idx, table in enumerate(doc.tables, 1):
            for r_idx, row in enumerate(table.rows, 1):
                cells = [cell.text.strip() for cell in row.cells]
                row_text = "\t".join(cells).strip()
                if row_text:
                    tables.append({
                        "id": f"T{t_idx:02d}R{r_idx:03d}",
                        "table_index": t_idx - 1,
                        "row_index": r_idx - 1,
                        "text": row_text,
                    })

    return doc, paragraphs, tables



def build_full_text(paragraphs: List[Dict], tables: List[Dict], include_tables: bool = True) -> str:
    body_lines = [f"[{p['id']}] {p['text']}" for p in paragraphs if p["text"]]
    if include_tables and tables:
        body_lines.append("\n【表格内容】")
        body_lines.extend([f"[{r['id']}] {r['text']}" for r in tables if r["text"]])
    return "\n".join(body_lines)



def _score_summary_paragraph(p: Dict) -> int:
    text = p["text"]
    if not text:
        return -10
    score = 0
    style = (p.get("style") or "").lower()
    if "heading" in style or "标题" in style:
        score += 8
    if re.match(r"^[一二三四五六七八九十0-9]+[、.）)]", text):
        score += 6
    if len(text) >= 20:
        score += 2
    if any(k in text for k in ["摘要", "总结", "结论", "建议", "背景", "目标", "方法", "结果", "问题", "风险"]):
        score += 5
    if re.search(r"\d", text):
        score += 2
    if text.startswith("•") or text.startswith("-"):
        score += 3
    return score



def _score_trend_paragraph(text: str) -> int:
    if isinstance(text, dict):
        text = text.get("text", "")
    if text is None:
        text = ""
    text = str(text)
    if not text:
        return -10
    score = 0
    digits = len(re.findall(r"\d", text))
    score += min(digits, 10)
    if re.search(r"\d{4}年|\d+月|\d+日|同比|环比|增长|下降|提升|减少|波动|峰值|低谷|占比|%|％", text):
        score += 8
    if "\t" in text:
        score += 5
    if any(k in text for k in ["趋势", "变化", "原因", "异常", "波动"]):
        score += 4
    return score



def distill_text(paragraphs: List[Dict], tables: List[Dict], mode: str, max_chars: int) -> str:
    """
    大文档快速蒸馏：保留标题、关键段、数字段、列表项和表格行，减少 LLM 输入量。
    """
    selected = []
    if mode == "summary":
        scored = sorted(paragraphs, key=_score_summary_paragraph, reverse=True)
        for p in scored:
            if p["text"]:
                selected.append(f"[{p['id']}] {p['text']}")
            if sum(len(x) for x in selected) >= max_chars:
                break
        # 保底：前后文各取一部分，避免只剩散点
        prefix = [f"[{p['id']}] {p['text']}" for p in paragraphs[:15] if p["text"]]
        suffix = [f"[{p['id']}] {p['text']}" for p in paragraphs[-15:] if p["text"]]
        selected = prefix + selected + suffix
    else:
        scored_p = sorted(paragraphs, key=lambda p: _score_trend_paragraph((p or {}).get('text', '')), reverse=True)
        for p in scored_p:
            if p["text"]:
                selected.append(f"[{p['id']}] {p['text']}")
            if sum(len(x) for x in selected) >= int(max_chars * 0.6):
                break
        table_lines = [f"[{r['id']}] {r['text']}" for r in tables if r['text']]
        for line in table_lines:
            selected.append(line)
            if sum(len(x) for x in selected) >= max_chars:
                break

    # 去重但保留顺序
    seen = set()
    uniq = []
    for item in selected:
        if item not in seen:
            uniq.append(item)
            seen.add(item)

    text = "\n".join(uniq)
    return text[:max_chars]



def chunk_text(text: str, max_chars: int = 5000) -> List[str]:
    if not text:
        return []
    lines = text.splitlines()
    chunks = []
    cur = []
    cur_len = 0
    for line in lines:
        line_len = len(line) + 1
        if cur and cur_len + line_len > max_chars:
            chunks.append("\n".join(cur))
            cur = [line]
            cur_len = line_len
        else:
            cur.append(line)
            cur_len += line_len
    if cur:
        chunks.append("\n".join(cur))
    return chunks



def estimate_char_budget(llm_kwargs: Dict, default_budget: int = 5000) -> int:
    try:
        from request_llms.bridge_all import model_info
        max_token = model_info[llm_kwargs['llm_model']]['max_token']
        return max(2500, min(int(max_token * 0.45), 12000))
    except Exception:
        return default_budget



def paragraph_issue_heuristics(paragraphs: List[Dict], max_candidates: int = 30) -> List[Dict]:
    """规则先筛一遍，把最可能有问题的段落交给大模型。"""
    informal_words = ["其实", "然后", "比较", "很多", "一些", "这个", "那个", "等等", "非常", "特别", "很", "较为"]
    candidates = []
    for p in paragraphs:
        text = p["text"]
        if not text or len(text) < 20:
            continue
        score = 0
        reasons = []

        if len(text) > 120:
            score += 4
            reasons.append("句段较长")
        if len(re.findall(r"[，、；：]", text)) >= 6:
            score += 3
            reasons.append("停顿过多")
        if re.search(r"([。！？；，])\1{1,}", text):
            score += 4
            reasons.append("标点重复")
        if re.search(r"\s{2,}", text):
            score += 2
            reasons.append("空格不规范")
        if re.search(r"[A-Za-z][\u4e00-\u9fff]|[\u4e00-\u9fff][A-Za-z]", text):
            score += 2
            reasons.append("中英文混排可能不规范")
        weak_hits = sum(w in text for w in informal_words)
        if weak_hits >= 2:
            score += 3
            reasons.append("表达偏口语化")
        if text.count("的") >= 8:
            score += 2
            reasons.append("表述可能拖沓")
        if "(" in text or "（" in text:
            if text.count("(") + text.count("（") >= 2:
                score += 2
                reasons.append("插入说明较多")
        if re.search(r"\d+\s+(kg|m|cm|%|℃)", text, flags=re.I):
            score += 1
            reasons.append("单位前后空格需检查")

        if score > 0:
            candidates.append({
                "paragraph_id": p["id"],
                "score": score,
                "reasons": reasons,
                "excerpt": text[:180],
                "text": text,
            })

    candidates.sort(key=lambda x: (-x["score"], x["paragraph_id"]))
    return candidates[:max_candidates]



def save_json(data: Dict, save_path: str):
    with open(save_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)



def mark_docx_paragraphs(input_fp: str, issues: List[Dict], output_fp: str):
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import RGBColor

    doc = Document(input_fp)
    pid_to_para = {f"P{i+1:04d}": para for i, para in enumerate(doc.paragraphs)}

    for idx, issue in enumerate(issues, 1):
        pid = issue.get("paragraph_id")
        para = pid_to_para.get(pid)
        if para is None:
            continue
        marker = f"【问题{idx}:{issue.get('severity', '中')}】"
        if para.runs:
            para.runs[0].text = marker + para.runs[0].text
            para.runs[0].font.bold = True
            para.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        else:
            run = para.add_run(marker)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        for run in para.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # 在文末补一个索引表，便于快速跳转核对
    doc.add_page_break()
    doc.add_heading("问题定位索引", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "问题ID"
    hdr[1].text = "段落ID"
    hdr[2].text = "严重度"
    hdr[3].text = "原文摘录"

    for idx, issue in enumerate(issues, 1):
        row = table.add_row().cells
        row[0].text = f"问题{idx}"
        row[1].text = issue.get("paragraph_id", "")
        row[2].text = issue.get("severity", "")
        row[3].text = issue.get("original_excerpt", issue.get("excerpt", ""))[:160]

    Path(os.path.dirname(output_fp) or ".").mkdir(parents=True, exist_ok=True)
    doc.save(output_fp)



def add_bullets(doc, title: str, items: List[str], level: int = 2):
    if not items:
        return
    doc.add_heading(title, level=level)
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")



def add_key_value_table(doc, title: str, rows: List[Tuple[str, str]], level: int = 2):
    if not rows:
        return
    doc.add_heading(title, level=level)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "字段"
    hdr[1].text = "内容"
    for k, v in rows:
        r = table.add_row().cells
        r[0].text = str(k)
        r[1].text = str(v)



def ensure_dir(path: str):
    Path(path).mkdir(parents=True, exist_ok=True)

