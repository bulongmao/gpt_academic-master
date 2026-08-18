from toolbox import update_ui
from toolbox import CatchException, report_exception
from toolbox import write_history_to_file, promote_file_to_downloadzone
from crazy_functions.crazy_utils import request_gpt_model_in_new_thread_with_ui_alive

import os
import sys
import re
import copy

_CURRENT_DIR = os.path.dirname(os.path.abspath(__file__)) if "__file__" in globals() else os.getcwd()
if _CURRENT_DIR not in sys.path:
    sys.path.insert(0, _CURRENT_DIR)

try:
    from crazy_functions.word_plugin_utils import (
        strip_model_noise,
        safe_load_json,
        now_str,
        sanitize_filename,
        read_docx_content,
        build_full_text,
        distill_text,
        chunk_text,
        estimate_char_budget,
    )
except ImportError:
    from word_plugin_utils import (
        strip_model_noise,
        safe_load_json,
        now_str,
        sanitize_filename,
        read_docx_content,
        build_full_text,
        distill_text,
        chunk_text,
        estimate_char_budget,
    )

fast_debug = False

SUMMARY_SECTION_TITLES = [("main_content_summary", "主要内容概括"), ("identified_issues_and_challenges", "识别问题与挑战"), ("evaluation_suggestions", "评估建议部分"), ("summary_conclusion", "总结结论")]


def _strip_protocol_noise(text: str) -> str:
    text = text or ""
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.S | re.I)
    text = re.sub(r"</?json>", "", text, flags=re.I)
    text = re.sub(r"</?answer>", "", text, flags=re.I)
    text = re.sub(r"</?output>", "", text, flags=re.I)
    text = re.sub(r"```(?:json|text|markdown)?\s*", "", text, flags=re.I)
    text = text.replace("```", "")
    text = re.sub(r"^\s*json\s*", "", text, flags=re.I)
    return text.strip()


def _clean_obj(obj):
    if isinstance(obj, dict):
        return {k: _clean_obj(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_clean_obj(v) for v in obj]
    if isinstance(obj, str):
        return _strip_protocol_noise(strip_model_noise(obj)).strip(" \n\r\t,")
    return obj


def _extract_json_field_text(text: str, keys):
    raw = text or ""
    for key in keys:
        m = re.search(rf'"{re.escape(key)}"\s*:\s*"((?:\\.|[^"\\])*)"', raw, flags=re.S)
        if m:
            s = m.group(1)
            s = s.replace('\\n', '\n').replace('\\t', '\t').replace('\\"', '"')
            return _strip_protocol_noise(s).strip()
    return ""


def _best_effort_text(text: str, preferred_keys, max_len: int) -> str:
    obj = safe_load_json(text, None)
    if isinstance(obj, dict):
        obj = _clean_obj(obj)
        for key in preferred_keys:
            value = obj.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()[:max_len]

    field_text = _extract_json_field_text(text, preferred_keys)
    if field_text:
        return field_text[:max_len]

    cleaned = _strip_protocol_noise(strip_model_noise(text))
    lines = []
    for line in cleaned.splitlines():
        s = line.strip().strip(",")
        if not s:
            continue
        if s in ["{", "}", "[", "]"]:
            continue
        if re.match(r'^"[^\"]+"\s*:', s):
            continue
        lines.append(s)
    merged = " ".join(lines).strip()
    return merged[:max_len]


def _safe_json_data(text: str, default: dict, preferred_keys, max_len: int) -> dict:
    data = safe_load_json(text, None)
    if not isinstance(data, dict):
        data = copy.deepcopy(default)
    else:
        data = _clean_obj(data)
        merged = copy.deepcopy(default)
        merged.update(data)
        data = merged

    fallback = _best_effort_text(text, preferred_keys=preferred_keys, max_len=max_len)
    for key in preferred_keys:
        if not str(data.get(key, "")).strip() and fallback:
            data[key] = fallback
    return data





def _ensure_list(value):
    if value is None:
        return []
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    if isinstance(value, str):
        txt = _strip_protocol_noise(strip_model_noise(value)).strip()
        if not txt:
            return []
        parts = re.split(r'[\n；;]+', txt)
        return [p.strip(' -•') for p in parts if p.strip(' -•')]
    return [str(value).strip()]


def _clean_section_item(text: str) -> str:
    text = _strip_protocol_noise(strip_model_noise(text or "")).strip()
    text = text.replace("•", "")
    text = re.sub(r'^\s*[-*]\s*', '', text)
    text = re.sub(r'^\s*\d+[\.、]\s*', '', text)
    text = re.sub(r'^\s*[一二三四五六七八九十]+、\s*', '', text)
    text = re.sub(r'^\s*#{1,6}\s*', '', text)
    text = re.sub(r'^\s*\*\*(.*?)\*\*\s*$', r'\1', text)
    text = re.sub(r'^\s*(顾问总结|总结与分析|执行摘要)\s*[：:]*\s*', '', text)
    text = text.strip("：: ")
    return text.strip()


def _join_section_source_text(raw_sections: dict, fallback_text: str, title_mapping: list) -> str:
    parts = []
    if fallback_text and str(fallback_text).strip():
        parts.append(str(fallback_text).strip())
    if isinstance(raw_sections, dict):
        for key, _ in title_mapping:
            value = raw_sections.get(key)
            if isinstance(value, list):
                joined = "\n".join(str(v) for v in value if str(v).strip())
                if joined.strip():
                    parts.append(joined.strip())
            elif isinstance(value, str) and value.strip():
                parts.append(value.strip())
    return "\n".join(parts).strip()


def _split_section_fallback_text(text: str, title_mapping: list) -> dict:
    """
    当模型没有正确返回 executive_summary_sections 时，
    尝试从 fallback_text 或 raw_sections 中按“1.主要内容概括 / 2.识别问题与挑战 / 3.评估建议部分 / 4.总结结论”反拆。
    """
    text = _strip_protocol_noise(strip_model_noise(text or ""))
    if not text.strip():
        return {k: [] for k, _ in title_mapping}

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("•", "\n• ")
    text = re.sub(r'\s*(?=###)', '\n', text)
    text = re.sub(r'\s*(?=(?:\d+[\.、]\s*\*{0,2}(?:主要内容概括|识别问题与挑战|评估建议部分|总结结论)\*{0,2}\s*[：:]))', '\n', text)
    text = re.sub(r'\s*(?=(?:\d+[\.、]\s*\*{0,2}(?:学科建设|人才引进|产学研合作|面临的挑战与改进建议|政策执行情况|未来改进方向|显著成果|面临挑战)\*{0,2}\s*[：:]))', '\n', text)
    text = re.sub(r'\n{2,}', '\n', text)

    for idx, (key, title) in enumerate(title_mapping, 1):
        pattern = rf'\s*{idx}\s*[\.、]?\s*\*{{0,2}}\s*{re.escape(title)}\s*\*{{0,2}}\s*[：:]'
        text = re.sub(pattern, f"\n@@SECTION::{key}@@\n", text, flags=re.I)

    result = {k: [] for k, _ in title_mapping}
    current_key = None

    for raw in text.split("\n"):
        line = raw.strip()
        if not line:
            continue

        if line.startswith("@@SECTION::") and line.endswith("@@"):
            current_key = line[len("@@SECTION::"):-2]
            continue

        matched = False
        for idx, (key, title) in enumerate(title_mapping, 1):
            m = re.match(
                rf'^\s*{idx}\s*[\.、]?\s*\*{{0,2}}\s*{re.escape(title)}\s*\*{{0,2}}\s*[：:]?\s*(.*)$',
                line,
                flags=re.I
            )
            if m:
                current_key = key
                rest = _clean_section_item(m.group(1))
                if rest:
                    result[current_key].append(rest)
                matched = True
                break
        if matched:
            continue

        if current_key:
            cleaned = _clean_section_item(line)
            if cleaned:
                result[current_key].append(cleaned)

    for key in result:
        cleaned_items = []
        seen = set()
        for item in result[key]:
            item = item.strip()
            if not item:
                continue
            if item in ["顾问总结", "总结与分析", "执行摘要", "以下是总结与分析"]:
                continue
            if item not in seen:
                cleaned_items.append(item)
                seen.add(item)
        result[key] = cleaned_items

    return result


def _build_sectioned_summary(data: dict, title_mapping: list, fallback_text: str = "") -> dict:
    raw_sections = data.get("executive_summary_sections") or data.get("summary_sections") or {}
    if not isinstance(raw_sections, dict):
        raw_sections = {}

    result = {}
    for key, _ in title_mapping:
        result[key] = _ensure_list(raw_sections.get(key))

    non_empty_cnt = sum(1 for v in result.values() if v)
    if non_empty_cnt <= 1:
        combined_text = _join_section_source_text(raw_sections, fallback_text, title_mapping)
        parsed = _split_section_fallback_text(combined_text, title_mapping)
        for key, _ in title_mapping:
            if parsed.get(key):
                result[key] = parsed[key]

    non_empty_cnt = sum(1 for v in result.values() if v)
    if non_empty_cnt == 0:
        fallback = _strip_protocol_noise(strip_model_noise(fallback_text or "")).strip()
        if fallback:
            result[title_mapping[0][0]] = [fallback]
    return result


def _split_long_section_item(text: str):
    """把单个很长的 section 条目尽量拆成多个小段/小点。"""
    text = _strip_protocol_noise(strip_model_noise(text or "")).strip()
    if not text:
        return []

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # 先把常见 markdown / 编号模式打散成换行
    text = re.sub(r'\s*(?=#{1,6}\s*)', '\n', text)
    text = re.sub(r'\s*(?=[-*•]\s+)', '\n', text)
    text = re.sub(r'\s*(?=\d+\s*[\.、）)]\s*)', '\n', text)
    text = re.sub(r'\s*(?=[一二三四五六七八九十]+[、）)])', '\n', text)
    text = re.sub(r'([：:])\s*(?=\d+\s*[\.、）)]\s*)', r'\1\n', text)
    text = re.sub(r'([：:])\s*(?=[一二三四五六七八九十]+[、）)])', r'\1\n', text)
    text = re.sub(r'\n{2,}', '\n', text).strip()

    parts = []
    for raw in text.split("\n"):
        seg = _clean_section_item(raw)
        if not seg:
            continue
        if seg in ["顾问总结", "总结与分析", "执行摘要", "以下是总结与分析"]:
            continue
        parts.append(seg)

    cleaned = []
    seen = set()
    for p in parts:
        if p not in seen:
            cleaned.append(p)
            seen.add(p)
    return cleaned


def _render_sectioned_summary(doc, section_data: dict, title_mapping: list):
    # 只显示有内容的 section；避免空 section 影响观感
    visible_sections = []
    for key, title in title_mapping:
        raw_items = _ensure_list(section_data.get(key))
        split_items = []
        for item in raw_items:
            split_items.extend(_split_long_section_item(item))
        if split_items:
            visible_sections.append((title, split_items))

    if not visible_sections:
        return

    for idx, (title, items) in enumerate(visible_sections, 1):
        p = doc.add_paragraph()
        p.add_run(f"{idx}.").bold = True
        p.add_run(title).bold = True
        p.add_run("：").bold = True
        for item in items:
            doc.add_paragraph(str(item), style="List Bullet")

def _add_inline_bold(paragraph, text: str):
    parts = re.split(r'(\*\*.*?\*\*)', text or "")
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _render_clean_block(doc, text: str):
    """
    将模型输出中的 markdown/伪结构文本尽量转成清晰的 Word 段落。
    """
    text = _strip_protocol_noise(strip_model_noise(text))
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r'^\s*以下是.*?[：:]\s*', '', text, flags=re.M)
    text = re.sub(r'\n{3,}', '\n\n', text).strip()

    if not text:
        doc.add_paragraph("")
        return

    for raw in text.split("\n"):
        line = raw.strip()
        if not line:
            continue

        m = re.match(r'^\s*#{1,6}\s*(.+?)\s*$', line)
        if m:
            p = doc.add_paragraph()
            run = p.add_run(m.group(1).strip())
            run.bold = True
            continue

        m = re.match(r'^\s*(\d+[\.、])\s*(.+?)\s*$', line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_inline_bold(p, m.group(2).strip())
            continue

        m = re.match(r'^\s*([一二三四五六七八九十]+、)\s*(.+?)\s*$', line)
        if m:
            p = doc.add_paragraph()
            run = p.add_run(f"{m.group(1)} {m.group(2).strip()}")
            run.bold = True
            continue

        m = re.match(r'^\s*[-*•]\s+(.+?)\s*$', line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_bold(p, m.group(1).strip())
            continue

        p = doc.add_paragraph()
        _add_inline_bold(p, line)


def export_summary_report(doc_title, processing_mode, raw_length, analysis_length, chunk_results, final_data, save_dir="./tmp"):
    import os
    import time
    from pathlib import Path
    from docx import Document

    Path(save_dir).mkdir(parents=True, exist_ok=True)
    out_path = os.path.join(save_dir, f"摘要报告_{sanitize_filename(doc_title)}_{int(time.time())}.docx")

    doc = Document()
    doc.add_heading(f"文本摘要报告 - {doc_title}", level=0)
    doc.add_paragraph(f"生成时间：{now_str()}")
    doc.add_paragraph(f"处理模式：{processing_mode}")
    doc.add_paragraph(f"原始文本长度：{raw_length} 字符")
    doc.add_paragraph(f"送入模型的分析文本长度：{analysis_length} 字符")

    doc.add_heading("执行摘要", level=1)
    section_data = _build_sectioned_summary(final_data, SUMMARY_SECTION_TITLES, final_data.get("executive_summary", ""))
    _render_sectioned_summary(doc, section_data, SUMMARY_SECTION_TITLES)

    conclusions = final_data.get("core_conclusions", [])
    if conclusions:
        doc.add_heading("核心结论", level=1)
        for item in conclusions:
            doc.add_paragraph(str(item), style="List Bullet")

    outline = final_data.get("document_outline", [])
    if outline:
        doc.add_heading("主题概览", level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "主题"
        hdr[1].text = "摘要"
        for row in outline:
            r = table.add_row().cells
            r[0].text = str(row.get("topic", ""))
            r[1].text = str(row.get("summary", ""))

    keywords = final_data.get("keywords", [])
    if keywords:
        doc.add_heading("关键词", level=1)
        doc.add_paragraph("、".join(map(str, keywords)))

    if chunk_results:
        doc.add_heading("分段摘要", level=1)
        for idx, item in enumerate(chunk_results, 1):
            doc.add_heading(f"片段 {idx}", level=2)
            doc.add_paragraph(item.get("chunk_summary", ""))
            points = item.get("key_points", [])
            if points:
                for p in points:
                    doc.add_paragraph(str(p), style="List Bullet")

    doc.save(out_path)
    return out_path


def _build_chunk_prompt(rel_path, chunk_idx, chunk_total, chunk_text):
    return (
        "你是一名中文政策与管理文档摘要助手。请从事实层面进行归纳，不要输出思考过程，不要解释格式。\n"
        "请严格仅输出 <json>...</json>，不要输出除 JSON 之外的任何前后缀。\n"
        "输出格式：\n"
        "<json>{\n"
        '  "chunk_summary": "180到260字的片段摘要，需交代该片段讨论主题、主要问题、主要措施或结论",\n'
        '  "key_points": ["要点1", "要点2", "要点3", "要点4"],\n'
        '  "keywords": ["关键词1", "关键词2", "关键词3"]\n'
        "}</json>\n\n"
        f"文件：{rel_path}；片段：{chunk_idx}/{chunk_total}\n"
        "待分析内容如下：\n"
        f"```text\n{chunk_text}\n```"
    )



def _build_merge_prompt(doc_name, chunk_json_list):
    import json
    return f"""你是一名高级中文编辑，负责将片段摘要整合为正式报告中的执行摘要。不要输出思考过程，不要解释格式，不要使用空泛套话。
请严格仅输出 <json>...</json>。
执行摘要必须严格按照以下四部分组织：
1. 主要内容概括；
2. 识别问题与挑战；
3. 评估建议部分；
4. 总结结论。
要求如下：
1. 四个部分都必须填写，任何一个部分都不能为空；
2. 每一部分都输出为要点列表，每部分 2 到 4 条；
3. 内容必须贴合原文，不能编造；
4. 不要把四个部分的内容混在同一个字段里；
5. 语言要正式、概括、便于写入报告；
6. 不要输出 Markdown 标记，不要输出 ###、**、-、1. 这种排版符号，由程序负责排版。
输出格式：
<json>{{
  "executive_summary_sections": {{
    "main_content_summary": ["要点1", "要点2"],
    "identified_issues_and_challenges": ["要点1", "要点2"],
    "evaluation_suggestions": ["要点1", "要点2"],
    "summary_conclusion": ["要点1", "要点2"]
  }},
  "executive_summary": "可选：用于兜底的简短自然段摘要",
  "core_conclusions": ["结论1", "结论2", "结论3", "结论4", "结论5"],
  "document_outline": [{{"topic": "主题1", "summary": "一句话概括"}}],
  "keywords": ["关键词1", "关键词2", "关键词3", "关键词4"]
}}</json>

文档名称：{doc_name}
片段摘要数据如下：
```json
{json.dumps(chunk_json_list, ensure_ascii=False, indent=2)}
```"""


def 解析docx(file_manifest, project_folder, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt):
    import os
    import json

    fast_threshold_chars = int(plugin_kwargs.get("fast_threshold_chars", 12000)) if isinstance(plugin_kwargs, dict) else 12000
    distill_max_chars = int(plugin_kwargs.get("distill_max_chars", 18000)) if isinstance(plugin_kwargs, dict) else 18000

    for fp in file_manifest:
        if not fp.lower().endswith(".docx"):
            raise RuntimeError("优化版当前建议仅处理 .docx 文件；.doc 请先转换为 .docx。")

        doc, paragraphs, tables = read_docx_content(fp, include_tables=True)
        full_text = build_full_text(paragraphs, tables, include_tables=True)
        raw_length = len(full_text)
        char_budget = estimate_char_budget(llm_kwargs)

        if raw_length <= fast_threshold_chars:
            analysis_text = full_text
            processing_mode = "全文摘要模式"
        else:
            analysis_text = distill_text(
                paragraphs=paragraphs,
                tables=tables,
                mode="summary",
                max_chars=max(distill_max_chars, char_budget)
            )
            processing_mode = "快速蒸馏摘要模式"

        chunks = chunk_text(analysis_text, max_chars=char_budget)
        if not chunks:
            chunks = [analysis_text[:char_budget]]

        chunk_results = []
        rel_path = os.path.relpath(fp, project_folder)

        for i, chunk in enumerate(chunks, 1):
            i_say = _build_chunk_prompt(rel_path, i, len(chunks), chunk)
            i_say_show_user = f"正在摘要：{os.path.abspath(fp)} 第 {i}/{len(chunks)} 个分析片段"
            gpt_say = yield from request_gpt_model_in_new_thread_with_ui_alive(
                inputs=i_say,
                inputs_show_user=i_say_show_user,
                llm_kwargs=llm_kwargs,
                chatbot=chatbot,
                history=[],
                sys_prompt="总结文章，只输出结论。"
            )
            data = _safe_json_data(
                gpt_say,
                default={"chunk_summary": "", "key_points": [], "keywords": []},
                preferred_keys=["chunk_summary", "summary", "executive_summary"],
                max_len=320,
            )
            chunk_results.append(data)

            clean_view = (data.get("chunk_summary") or _best_effort_text(gpt_say, ["chunk_summary"], 220) or "（片段摘要解析失败）")
            chatbot[-1] = (i_say_show_user, clean_view)
            history.extend([i_say_show_user, clean_view])

        merge_prompt = _build_merge_prompt(os.path.basename(fp), chunk_results)
        merge_show = f"正在汇总最终摘要：{os.path.abspath(fp)}"
        gpt_final = yield from request_gpt_model_in_new_thread_with_ui_alive(
            inputs=merge_prompt,
            inputs_show_user=merge_show,
            llm_kwargs=llm_kwargs,
            chatbot=chatbot,
            history=[],
            sys_prompt="总结文章，只输出结论。"
        )
        final_data = _safe_json_data(
            gpt_final,
            default={"executive_summary": "", "executive_summary_sections": {}, "core_conclusions": [], "document_outline": [], "keywords": []},
            preferred_keys=["executive_summary", "summary", "chunk_summary"],
            max_len=1600,
        )
        final_text = final_data.get("executive_summary") or _best_effort_text(gpt_final, ["executive_summary", "summary"], 1200)
        final_data["executive_summary"] = final_text
        clean_final = final_text[:420] + ("……" if len(final_text) > 420 else "")
        chatbot[-1] = (merge_show, clean_final)
        history.extend([merge_show, clean_final])

        report_path = export_summary_report(
            doc_title=os.path.basename(fp),
            processing_mode=processing_mode,
            raw_length=raw_length,
            analysis_length=len(analysis_text),
            chunk_results=chunk_results,
            final_data=final_data,
            save_dir="./tmp"
        )
        json_path = report_path.replace(".docx", ".json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump({
                "file": fp,
                "processing_mode": processing_mode,
                "raw_length": raw_length,
                "analysis_length": len(analysis_text),
                "chunk_results": chunk_results,
                "final_data": final_data,
            }, f, ensure_ascii=False, indent=2)

        promote_file_to_downloadzone(report_path, chatbot=chatbot)
        promote_file_to_downloadzone(json_path, chatbot=chatbot)
        chatbot.append(("📄 已生成摘要报告", report_path))
        yield from update_ui(chatbot=chatbot, history=history)

    res = write_history_to_file(history)
    promote_file_to_downloadzone(res, chatbot=chatbot)
    chatbot.append(("全部摘要任务完成", res))
    yield from update_ui(chatbot=chatbot, history=history)


@CatchException
def 总结word文档(txt, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt, user_request):
    import glob
    import os

    chatbot.append([
        "函数插件功能？",
        "批量总结 Word 文档（优化版）：大文档自动进入快速蒸馏模式，只输出清爽结论，并生成结构化报告。"
    ])
    yield from update_ui(chatbot=chatbot, history=history)

    try:
        from docx import Document  # noqa: F401
    except Exception:
        report_exception(
            chatbot, history,
            a=f"解析项目: {txt}",
            b="导入依赖失败，请安装：```pip install --upgrade python-docx pywin32```。"
        )
        yield from update_ui(chatbot=chatbot, history=history)
        return

    history = []

    if os.path.exists(txt):
        project_folder = txt
    else:
        if txt == "":
            txt = '空空如也的输入栏'
        report_exception(chatbot, history, a=f"解析项目: {txt}", b=f"找不到本地项目或无权访问: {txt}")
        yield from update_ui(chatbot=chatbot, history=history)
        return

    if txt.endswith('.docx'):
        file_manifest = [txt]
        project_folder = os.path.dirname(os.path.abspath(txt)) or os.getcwd()
    else:
        project_folder = os.path.abspath(project_folder)
        file_manifest = [f for f in glob.glob(f'{project_folder}/**/*.docx', recursive=True)]

    if len(file_manifest) == 0:
        report_exception(chatbot, history, a=f"解析项目: {txt}", b=f"找不到任何 .docx 文件: {txt}")
        yield from update_ui(chatbot=chatbot, history=history)
        return

    yield from 解析docx(file_manifest, project_folder, llm_kwargs, plugin_kwargs or {}, chatbot, history, system_prompt)
