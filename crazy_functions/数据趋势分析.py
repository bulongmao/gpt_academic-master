from toolbox import update_ui
from toolbox import CatchException, report_exception
from toolbox import write_history_to_file, promote_file_to_downloadzone
from crazy_functions.crazy_utils import request_gpt_model_in_new_thread_with_ui_alive

import os
import sys
import re
import copy

_CURRENT_DIR = os.path.dirname(os.path.abspath(__file__)) if "__file__" in globals() else os.getcwd()
if _CURRENT_DIR not in sys.path:
    sys.path.insert(0, _CURRENT_DIR)

try:
    from crazy_functions.word_plugin_utils import (
        strip_model_noise,
        safe_load_json,
        now_str,
        sanitize_filename,
        read_docx_content,
        build_full_text,
        distill_text,
        chunk_text,
        estimate_char_budget,
    )
except ImportError:
    from word_plugin_utils import (
        strip_model_noise,
        safe_load_json,
        now_str,
        sanitize_filename,
        read_docx_content,
        build_full_text,
        distill_text,
        chunk_text,
        estimate_char_budget,
    )

fast_debug = False


def _strip_protocol_noise(text: str) -> str:
    text = text or ""
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.S | re.I)
    text = re.sub(r"</?json>", "", text, flags=re.I)
    text = re.sub(r"```(?:json|text|markdown)?\s*", "", text, flags=re.I)
    text = text.replace("```", "")
    return text.strip()


def _clean_obj(obj):
    if isinstance(obj, dict):
        return {k: _clean_obj(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_clean_obj(v) for v in obj]
    if isinstance(obj, str):
        return _strip_protocol_noise(strip_model_noise(obj)).strip(" \n\r\t,")
    return obj


def _extract_json_field_text(text: str, keys):
    raw = text or ""
    for key in keys:
        m = re.search(rf'"{re.escape(key)}"\s*:\s*"((?:\\.|[^"\\])*)"', raw, flags=re.S)
        if m:
            s = m.group(1)
            s = s.replace('\\n', '\n').replace('\\t', '\t').replace('\\"', '"')
            return _strip_protocol_noise(s).strip()
    return ""


def _best_effort_text(text: str, preferred_keys, max_len: int) -> str:
    obj = safe_load_json(text, None)
    if isinstance(obj, dict):
        obj = _clean_obj(obj)
        for key in preferred_keys:
            value = obj.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()[:max_len]
    field_text = _extract_json_field_text(text, preferred_keys)
    if field_text:
        return field_text[:max_len]
    cleaned = _strip_protocol_noise(strip_model_noise(text))
    lines = []
    for line in cleaned.splitlines():
        s = line.strip().strip(",")
        if not s or s in ["{", "}", "[", "]"]:
            continue
        if re.match(r'^"[^\"]+"\s*:', s):
            continue
        lines.append(s)
    return " ".join(lines).strip()[:max_len]


def _safe_json_data(text: str, default: dict, preferred_keys, max_len: int) -> dict:
    data = safe_load_json(text, None)
    if not isinstance(data, dict):
        data = copy.deepcopy(default)
    else:
        data = _clean_obj(data)
        merged = copy.deepcopy(default)
        merged.update(data)
        data = merged
    fallback = _best_effort_text(text, preferred_keys=preferred_keys, max_len=max_len)
    for key in preferred_keys:
        if not str(data.get(key, "")).strip() and fallback:
            data[key] = fallback
    return data



def _add_inline_bold(paragraph, text: str):
    parts = re.split(r'(\*\*.*?\*\*)', text or "")
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _render_clean_block(doc, text: str):
    """
    将模型输出中的 markdown/伪结构文本尽量转成清晰的 Word 段落。
    """
    text = _strip_protocol_noise(strip_model_noise(text))
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r'^\s*以下是.*?[：:]\s*', '', text, flags=re.M)
    text = re.sub(r'\n{3,}', '\n\n', text).strip()

    if not text:
        doc.add_paragraph("")
        return

    for raw in text.split("\n"):
        line = raw.strip()
        if not line:
            continue

        m = re.match(r'^\s*#{1,6}\s*(.+?)\s*$', line)
        if m:
            p = doc.add_paragraph()
            run = p.add_run(m.group(1).strip())
            run.bold = True
            continue

        m = re.match(r'^\s*(\d+[\.、])\s*(.+?)\s*$', line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_inline_bold(p, m.group(2).strip())
            continue

        m = re.match(r'^\s*([一二三四五六七八九十]+、)\s*(.+?)\s*$', line)
        if m:
            p = doc.add_paragraph()
            run = p.add_run(f"{m.group(1)} {m.group(2).strip()}")
            run.bold = True
            continue

        m = re.match(r'^\s*[-*•]\s+(.+?)\s*$', line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_bold(p, m.group(1).strip())
            continue

        p = doc.add_paragraph()
        _add_inline_bold(p, line)


def export_trend_report(doc_title, processing_mode, raw_length, analysis_length, chunk_results, final_data, save_dir="./tmp"):
    import os
    import time
    from pathlib import Path
    from docx import Document

    Path(save_dir).mkdir(parents=True, exist_ok=True)
    out_path = os.path.join(save_dir, f"趋势分析报告_{sanitize_filename(doc_title)}_{int(time.time())}.docx")

    doc = Document()
    doc.add_heading(f"数据趋势分析报告 - {doc_title}", level=0)
    doc.add_paragraph(f"生成时间：{now_str()}")
    doc.add_paragraph(f"处理模式：{processing_mode}")
    doc.add_paragraph(f"原始文本长度：{raw_length} 字符")
    doc.add_paragraph(f"送入模型的分析文本长度：{analysis_length} 字符")

    doc.add_heading("总体结论", level=1)
    _render_clean_block(doc, final_data.get("overall_conclusion", ""))

    trends = final_data.get("main_trends", [])
    if trends:
        doc.add_heading("主要趋势", level=1)
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "指标/主题"
        hdr[1].text = "趋势方向"
        hdr[2].text = "证据"
        hdr[3].text = "可能原因"
        for item in trends:
            row = table.add_row().cells
            row[0].text = str(item.get("metric", ""))
            row[1].text = str(item.get("direction", ""))
            row[2].text = str(item.get("evidence", ""))
            row[3].text = str(item.get("reason", ""))

    anomalies = final_data.get("anomalies", [])
    if anomalies:
        doc.add_heading("异常点", level=1)
        for item in anomalies:
            doc.add_paragraph(str(item), style="List Bullet")

    suggestions = final_data.get("management_suggestions", [])
    if suggestions:
        doc.add_heading("建议", level=1)
        for item in suggestions:
            doc.add_paragraph(str(item), style="List Bullet")

    if chunk_results:
        doc.add_heading("分段分析记录", level=1)
        for idx, item in enumerate(chunk_results, 1):
            doc.add_heading(f"片段 {idx}", level=2)
            doc.add_paragraph(item.get("segment_summary", ""))
            for t in item.get("main_trends", []):
                doc.add_paragraph(
                    f"{t.get('metric', '')}：{t.get('direction', '')}；证据：{t.get('evidence', '')}",
                    style="List Bullet"
                )

    doc.save(out_path)
    return out_path



def _build_chunk_prompt(rel_path, chunk_idx, chunk_total, chunk_text):
    return (
        "你是一名数据分析助手。请基于文本中的数据、时间、比例和描述，识别趋势、异常和可能原因。"
        "不要输出思考过程，不要写长篇铺垫。\n"
        "请严格仅输出 <json>...</json>。\n"
        "输出格式：\n"
        "<json>{\n"
        '  "segment_summary": "180到260字的片段概括，需要交代主要变化方向、关键证据与异常情况",\n'
        '  "main_trends": [{"metric": "指标/主题", "direction": "上升/下降/波动/持平", "evidence": "原文证据", "reason": "可能原因"}],\n'
        '  "anomalies": ["异常点1", "异常点2"]\n'
        "}</json>\n\n"
        f"文件：{rel_path}；片段：{chunk_idx}/{chunk_total}\n"
        "待分析内容如下：\n"
        f"```text\n{chunk_text}\n```"
    )



def _build_merge_prompt(doc_name, chunk_json_list):
    import json
    return (
        "你是一名管理分析顾问。请把下面的片段趋势结果整合成正式报告中的总体趋势结论，"
        "优先保留最重要的趋势、异常、原因和建议，不要输出思考过程。\n"
        "请严格仅输出 <json>...</json>。\n"
        "要求如下：\n"
        "1. overall_conclusion 必须写成 700 到 1000 字的完整长摘要；\n"
        "2. 需要覆盖：分析对象、总体变化、重要增减方向、关键异常、原因判断、管理含义、后续建议；\n"
        "3. 尽量用报告语言写成连贯段落，不要只做列表拼接。\n"
        "输出格式：\n"
        "<json>{\n"
        '  "overall_conclusion": "700到1000字的总体趋势结论",\n'
        '  "main_trends": [{"metric": "指标/主题", "direction": "上升/下降/波动/持平", "evidence": "最关键证据", "reason": "可能原因"}],\n'
        '  "anomalies": ["异常点1", "异常点2"],\n'
        '  "management_suggestions": ["建议1", "建议2", "建议3"]\n'
        "}</json>\n\n"
        f"文档名称：{doc_name}\n"
        "片段结果如下：\n"
        f"```json\n{json.dumps(chunk_json_list, ensure_ascii=False, indent=2)}\n```"
    )



def 解析word趋势(file_manifest, project_folder, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt):
    import os
    import json

    fast_threshold_chars = int(plugin_kwargs.get("fast_threshold_chars", 15000)) if isinstance(plugin_kwargs, dict) else 15000
    distill_max_chars = int(plugin_kwargs.get("distill_max_chars", 22000)) if isinstance(plugin_kwargs, dict) else 22000

    for fp in file_manifest:
        if not fp.lower().endswith(".docx"):
            raise RuntimeError("优化版当前建议仅处理 .docx 文件；.doc 请先转换为 .docx。")

        doc, paragraphs, tables = read_docx_content(fp, include_tables=True)
        full_text = build_full_text(paragraphs, tables, include_tables=True)
        raw_length = len(full_text)
        char_budget = estimate_char_budget(llm_kwargs)

        if raw_length <= fast_threshold_chars:
            analysis_text = full_text
            processing_mode = "全文趋势分析模式"
        else:
            analysis_text = distill_text(
                paragraphs=paragraphs,
                tables=tables,
                mode="trend",
                max_chars=max(distill_max_chars, char_budget)
            )
            processing_mode = "快速蒸馏趋势分析模式"

        chunks = chunk_text(analysis_text, max_chars=char_budget)
        if not chunks:
            chunks = [analysis_text[:char_budget]]

        chunk_results = []
        rel_path = os.path.relpath(fp, project_folder)

        for i, chunk in enumerate(chunks, 1):
            i_say = _build_chunk_prompt(rel_path, i, len(chunks), chunk)
            i_say_show_user = f"正在分析趋势：{os.path.abspath(fp)} 第 {i}/{len(chunks)} 个分析片段"
            gpt_say = yield from request_gpt_model_in_new_thread_with_ui_alive(
                inputs=i_say,
                inputs_show_user=i_say_show_user,
                llm_kwargs=llm_kwargs,
                chatbot=chatbot,
                history=[],
                sys_prompt="数据趋势分析，只输出结论。"
            )
            data = _safe_json_data(
                gpt_say,
                default={"segment_summary": "", "main_trends": [], "anomalies": []},
                preferred_keys=["segment_summary", "summary", "overall_conclusion"],
                max_len=320,
            )
            chunk_results.append(data)

            clean_view = (data.get("segment_summary") or _best_effort_text(gpt_say, ["segment_summary"], 220) or "（片段趋势解析失败）")
            chatbot[-1] = (i_say_show_user, clean_view)
            history.extend([i_say_show_user, clean_view])

        merge_prompt = _build_merge_prompt(os.path.basename(fp), chunk_results)
        merge_show = f"正在汇总趋势结论：{os.path.abspath(fp)}"
        gpt_final = yield from request_gpt_model_in_new_thread_with_ui_alive(
            inputs=merge_prompt,
            inputs_show_user=merge_show,
            llm_kwargs=llm_kwargs,
            chatbot=chatbot,
            history=[],
            sys_prompt="数据趋势分析，只输出结论。"
        )
        final_data = _safe_json_data(
            gpt_final,
            default={"overall_conclusion": "", "main_trends": [], "anomalies": [], "management_suggestions": []},
            preferred_keys=["overall_conclusion", "summary", "segment_summary"],
            max_len=1400,
        )
        final_text = final_data.get("overall_conclusion") or _best_effort_text(gpt_final, ["overall_conclusion", "summary"], 1000)
        final_data["overall_conclusion"] = final_text
        clean_final = final_text[:420] + ("……" if len(final_text) > 420 else "")
        chatbot[-1] = (merge_show, clean_final)
        history.extend([merge_show, clean_final])

        report_path = export_trend_report(
            doc_title=os.path.basename(fp),
            processing_mode=processing_mode,
            raw_length=raw_length,
            analysis_length=len(analysis_text),
            chunk_results=chunk_results,
            final_data=final_data,
            save_dir="./tmp"
        )
        json_path = report_path.replace(".docx", ".json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump({
                "file": fp,
                "processing_mode": processing_mode,
                "raw_length": raw_length,
                "analysis_length": len(analysis_text),
                "chunk_results": chunk_results,
                "final_data": final_data,
            }, f, ensure_ascii=False, indent=2)

        promote_file_to_downloadzone(report_path, chatbot=chatbot)
        promote_file_to_downloadzone(json_path, chatbot=chatbot)
        chatbot.append(("📄 已生成趋势分析报告", report_path))
        yield from update_ui(chatbot=chatbot, history=history)

    res = write_history_to_file(history)
    promote_file_to_downloadzone(res, chatbot=chatbot)
    chatbot.append(("全部趋势分析任务完成", res))
    yield from update_ui(chatbot=chatbot, history=history)


@CatchException
def 数据趋势分析(txt, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt, user_request):
    import glob
    import os

    chatbot.append([
        "函数插件功能？",
        "批量读取 Word 文档并做趋势分析（优化版）：大文档优先抽取关键数字与表格，再输出简洁结论。"
    ])
    yield from update_ui(chatbot=chatbot, history=history)

    try:
        from docx import Document  # noqa: F401
    except Exception:
        report_exception(
            chatbot, history,
            a=f"解析项目: {txt}",
            b="缺少依赖，请安装：```pip install --upgrade python-docx pywin32```。"
        )
        yield from update_ui(chatbot=chatbot, history=history)
        return

    history = []

    if os.path.exists(txt):
        project_folder = txt
    else:
        if txt == "":
            txt = '空输入'
        report_exception(chatbot, history, a=f"解析项目: {txt}", b=f"找不到路径或无权限访问: {txt}")
        yield from update_ui(chatbot=chatbot, history=history)
        return

    if txt.endswith('.docx'):
        file_manifest = [txt]
        project_folder = os.path.dirname(os.path.abspath(txt)) or os.getcwd()
    else:
        project_folder = os.path.abspath(project_folder)
        file_manifest = [f for f in glob.glob(f'{project_folder}/**/*.docx', recursive=True)]

    if len(file_manifest) == 0:
        report_exception(chatbot, history, a=f"解析项目: {txt}", b="未找到任何 .docx 文件。")
        yield from update_ui(chatbot=chatbot, history=history)
        return

    yield from 解析word趋势(file_manifest, project_folder, plugin_kwargs=plugin_kwargs or {}, llm_kwargs=llm_kwargs,
                     chatbot=chatbot, history=history, system_prompt=system_prompt)
